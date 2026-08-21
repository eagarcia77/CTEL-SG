from __future__ import annotations
import io, re, json, html
from pathlib import Path
from typing import List, Dict, Tuple
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HEADING_1 = {
    'introducción','objetivos','objetivos del módulo','palabras clave',
    'lecturas y recursos requeridos','contenido','integración de conceptos',
    'conclusión','referencias'
}
YEAR_RE = re.compile(r'\b(19|20)\d{2}[a-z]?\b')
FIG_RE = re.compile(r'^\s*figura\s+(\d+)\b', re.I)
TABLE_RE = re.compile(r'^\s*tabla\s+(\d+)\b', re.I)
TOPIC_RE = re.compile(r'^\s*tema\s+\d+\b', re.I)
PAREN_CIT_RE = re.compile(r'\(([^()]*?(?:19|20)\d{2}[a-z]?[^()]*)\)')

def clean(text: str) -> str:
    return re.sub(r'\s+', ' ', text or '').strip()

def is_heading(text: str) -> Tuple[bool, int]:
    t = clean(text); low = t.lower()
    if low in HEADING_1 or TOPIC_RE.match(t): return True, 1
    if low in {'adquiere','refuerza'}: return True, 2
    if t and len(t) <= 120 and t.endswith(':'): return True, 2
    return False, 0

def is_fig(text): return bool(FIG_RE.match(clean(text)))
def is_table(text): return bool(TABLE_RE.match(clean(text)))
def is_note(text): return clean(text).lower().startswith('nota.')
def is_example(text): return clean(text).lower().startswith('ejemplo real.')

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge, edge_data in kwargs.items():
        tag = 'w:' + edge
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag); tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn('w:' + key), str(value))

def apply_apa_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'val':'nil'}, bottom={'val':'nil'}, left={'val':'nil'}, right={'val':'nil'}, insideH={'val':'nil'}, insideV={'val':'nil'})
    if not table.rows: return
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={'val':'single','sz':'8','color':'000000'}, bottom={'val':'single','sz':'8','color':'000000'}, left={'val':'nil'}, right={'val':'nil'})
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={'val':'single','sz':'8','color':'000000'}, left={'val':'nil'}, right={'val':'nil'})

def configure_docx(doc, font, font_size, profile):
    for sec in doc.sections:
        sec.top_margin = Inches(1); sec.bottom_margin = Inches(1); sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    normal = doc.styles['Normal']
    normal.font.name = font; normal._element.rPr.rFonts.set(qn('w:eastAsia'), font); normal.font.size = Pt(font_size)
    pf = normal.paragraph_format; pf.line_spacing = 2; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for name in ('APA7 Level 1','APA7 Level 2'):
        if name not in doc.styles:
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            s = doc.styles[name]
        s.font.name = font; s._element.rPr.rFonts.set(qn('w:eastAsia'), font); s.font.size = Pt(font_size); s.font.bold = True
        s.paragraph_format.line_spacing = 2; s.paragraph_format.space_before = Pt(0); s.paragraph_format.space_after = Pt(0); s.paragraph_format.keep_with_next = True
    doc.styles['APA7 Level 1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if profile == 'strict' else WD_ALIGN_PARAGRAPH.LEFT

def format_paragraphs(doc, profile):
    in_refs = False; pending_fig_title = False; pending_table_title = False
    for p in doc.paragraphs:
        text = clean(p.text)
        if not text: continue
        low = text.lower(); heading, level = is_heading(text)
        p.paragraph_format.line_spacing = 2; p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if low == 'referencias':
            in_refs = True; p.style = doc.styles['APA7 Level 1']; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Inches(0); continue
        if heading:
            p.style = doc.styles['APA7 Level 1' if level == 1 else 'APA7 Level 2']; p.paragraph_format.first_line_indent = Inches(0); in_refs = False; continue
        if in_refs:
            p.paragraph_format.first_line_indent = Inches(-0.5); p.paragraph_format.left_indent = Inches(0.5); continue
        if is_fig(text):
            for r in p.runs: r.bold=True; r.italic=False
            p.paragraph_format.first_line_indent = Inches(0); p.paragraph_format.keep_with_next = True; pending_fig_title=True; continue
        if pending_fig_title:
            for r in p.runs: r.italic=True; r.bold=False
            p.paragraph_format.first_line_indent = Inches(0); p.paragraph_format.keep_with_next = True; pending_fig_title=False; continue
        if is_table(text):
            for r in p.runs: r.bold=True; r.italic=False
            p.paragraph_format.first_line_indent = Inches(0); p.paragraph_format.keep_with_next = True; pending_table_title=True; continue
        if pending_table_title:
            for r in p.runs: r.italic=True; r.bold=False
            p.paragraph_format.first_line_indent = Inches(0); p.paragraph_format.keep_with_next = True; pending_table_title=False; continue
        if is_note(text):
            p.paragraph_format.first_line_indent = Inches(0)
            if p.runs and p.runs[0].text.strip().lower().startswith('nota.'): p.runs[0].italic=True
            continue
        if p.style and p.style.name and 'List' in p.style.name:
            p.paragraph_format.first_line_indent = Inches(0)
        else:
            p.paragraph_format.first_line_indent = Inches(0.5)
        if is_example(text) and p.runs and p.runs[0].text.strip().lower().startswith('ejemplo real.'):
            p.runs[0].bold = True

def format_tables(doc, font, font_size):
    for table in doc.tables:
        apply_apa_table_borders(table)
        for i,row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 2; p.paragraph_format.first_line_indent = Inches(0); p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        run.font.name = font; run._element.rPr.rFonts.set(qn('w:eastAsia'), font); run.font.size = Pt(font_size)
                        if i == 0: run.bold = True

def extract_pdf(data):
    reader = PdfReader(io.BytesIO(data)); blocks=[]
    for page in reader.pages:
        txt = page.extract_text() or ''
        for b in re.split(r'\n\s*\n', txt):
            b = clean(b.replace('\n',' '))
            if b: blocks.append(b)
    return blocks

def extract_html(data):
    soup = BeautifulSoup(data.decode('utf-8', errors='ignore'), 'lxml'); blocks=[]
    for el in soup.find_all(['h1','h2','h3','p','li','figcaption']):
        t = clean(el.get_text(' ', strip=True))
        if t: blocks.append(t)
    return blocks

def build_docx(blocks, output, font, font_size, profile):
    doc = Document(); configure_docx(doc,font,font_size,profile)
    for b in blocks: doc.add_paragraph(b)
    format_paragraphs(doc,profile); doc.save(output)

def docx_blocks(doc): return [clean(p.text) for p in doc.paragraphs if clean(p.text)]

def audit_text(blocks):
    warnings=[]; positives=[]
    ref_idx = next((i for i,x in enumerate(blocks) if clean(x).lower()=='referencias'), None)
    refs = blocks[ref_idx+1:] if ref_idx is not None else []; body = blocks[:ref_idx] if ref_idx is not None else blocks
    if ref_idx is None: warnings.append('No se detectó una sección titulada “Referencias”.')
    else: positives.append('Se detectó la sección de Referencias.')
    citations=[]
    for i,text in enumerate(body):
        for m in PAREN_CIT_RE.finditer(text): citations.append((i+1,m.group(1)))
    if citations: positives.append(f'Se detectaron {len(citations)} grupos de citas parentéticas.')
    else: warnings.append('No se detectaron citas parentéticas autor-fecha.')
    for para,content in citations:
        parts=[p.strip() for p in content.split(';') if YEAR_RE.search(p)]
        if len(parts)>1:
            keys=[re.sub(r'[^A-Za-zÁÉÍÓÚÑáéíóúñ]','',p.split(',')[0]).lower() for p in parts]
            if keys != sorted(keys): warnings.append(f'Posible orden no alfabético en una cita múltiple del bloque {para}: ({content}).')
    before=''; figures=set()
    for text in body:
        m=FIG_RE.match(clean(text))
        if m:
            n=m.group(1)
            if re.search(rf'\bFigura\s+{re.escape(n)}\b', before, flags=re.I): positives.append(f'Figura {n}: existe una mención previa en el texto.')
            else: warnings.append(f'Figura {n}: no se detectó una mención previa antes del rótulo.')
            figures.add(n)
        before += ' ' + text
    if refs:
        keys=[]
        for r in refs:
            key=re.split(r'[\.,(]', clean(r), 1)[0].strip().lower()
            if key: keys.append(key)
        if keys == sorted(keys): positives.append('Las referencias parecen estar en orden alfabético.')
        else: warnings.append('Las referencias pueden no estar completamente en orden alfabético.')
    lower={clean(x).lower() for x in blocks}
    for section in ('introducción','conclusión','referencias'):
        if section in lower: positives.append(f'Se detectó la sección “{section.title()}”.')
        else: warnings.append(f'No se detectó la sección “{section.title()}”.')
    return {'warnings':list(dict.fromkeys(warnings)),'positives':list(dict.fromkeys(positives)),'summary':{'blocks':len(blocks),'citations_detected':len(citations),'references_detected':len(refs),'figures_detected':len(figures)}}

def make_blackboard_html(blocks, profile, font, font_size):
    body=[]; in_refs=False; pending_fig=False; pending_table=False
    for text in blocks:
        esc=html.escape(text); low=clean(text).lower(); heading,level=is_heading(text)
        if low=='referencias': in_refs=True; body.append('<h2 class="apa-heading apa-center">Referencias</h2>'); continue
        if heading: in_refs=False; tag='h2' if level==1 else 'h3'; body.append(f'<{tag} class="apa-heading">{esc}</{tag}>'); continue
        if in_refs: body.append(f'<p class="apa-reference">{esc}</p>'); continue
        if is_fig(text): body.append(f'<p class="apa-figure-number">{esc}</p>'); pending_fig=True; continue
        if pending_fig: body.append(f'<p class="apa-figure-title">{esc}</p>'); pending_fig=False; continue
        if is_table(text): body.append(f'<p class="apa-table-number">{esc}</p>'); pending_table=True; continue
        if pending_table: body.append(f'<p class="apa-table-title">{esc}</p>'); pending_table=False; continue
        if is_note(text): body.append(f'<p class="apa-note">{esc}</p>'); continue
        body.append(f'<p class="apa-paragraph">{esc}</p>')
    align = 'center' if profile=='strict' else 'left'
    css=f'''<style>
.apa-module{{font-family:{font},Arial,Helvetica,sans-serif;font-size:{font_size}pt;line-height:2;color:#000;background:#fff;max-width:100%}}
.apa-paragraph{{margin:0;text-align:left;text-indent:.5in}} .apa-heading{{margin:0;font-size:{font_size}pt;line-height:2;font-weight:700;text-align:{align}}}.apa-center{{text-align:center}}
.apa-reference{{margin:0;padding-left:.5in;text-indent:-.5in;line-height:2}} .apa-figure-number,.apa-table-number{{margin:0;font-weight:700;text-indent:0;break-after:avoid-page;page-break-after:avoid}}
.apa-figure-title,.apa-table-title{{margin:0 0 8pt 0;font-style:italic;text-indent:0;break-after:avoid-page;page-break-after:avoid}} .apa-note{{margin:8pt 0 0 0;text-indent:0}} .apa-module img{{max-width:100%;height:auto}}
@media print{{@page{{size:letter;margin:1in}}.apa-module,.apa-module *{{overflow:visible!important;max-height:none!important}}.apa-module img,.apa-module figure{{break-inside:avoid-page;page-break-inside:avoid}}.apa-module table{{width:100%;border-collapse:collapse}}}}
</style>'''
    return css+'\n<div class="apa-module">\n'+'\n'.join(body)+'\n</div>\n'

def format_uploaded_module(filename,data,output_dir,profile,font,font_size):
    base=Path(filename).stem; suffix=Path(filename).suffix.lower(); output_dir.mkdir(exist_ok=True)
    safe=re.sub(r'[^A-Za-z0-9_-]+','_',base)[:80] or 'modulo'
    docx_out=output_dir/f'{safe}_APA7.docx'; html_out=output_dir/f'{safe}_Blackboard_APA7.html'; report_out=output_dir/f'{safe}_APA7_report.json'
    if suffix=='.docx':
        doc=Document(io.BytesIO(data)); configure_docx(doc,font,font_size,profile); format_paragraphs(doc,profile); format_tables(doc,font,font_size); blocks=docx_blocks(doc); doc.save(docx_out)
    elif suffix=='.pdf': blocks=extract_pdf(data); build_docx(blocks,docx_out,font,font_size,profile)
    elif suffix in {'.html','.htm'}: blocks=extract_html(data); build_docx(blocks,docx_out,font,font_size,profile)
    elif suffix=='.txt':
        text=data.decode('utf-8',errors='ignore'); blocks=[clean(x) for x in re.split(r'\n\s*\n|\n',text) if clean(x)]; build_docx(blocks,docx_out,font,font_size,profile)
    else: raise ValueError('Formato no compatible.')
    audit=audit_text(blocks); audit.update({'profile':profile,'font':font,'font_size':font_size,'source_file':filename})
    html_out.write_text(make_blackboard_html(blocks,profile,font,font_size),encoding='utf-8'); report_out.write_text(json.dumps(audit,ensure_ascii=False,indent=2),encoding='utf-8')
    return {'source':filename,'docx':docx_out.name,'html':html_out.name,'report':report_out.name,'audit':audit,'note':'En DOCX se conservan las imágenes existentes al reestilizar. En PDF/HTML/TXT se reconstruye el texto y puede perderse estructura visual compleja.'}
